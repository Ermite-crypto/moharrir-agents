from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

MAX_FILE_BYTES = 10 * 1024 * 1024
MAX_TOTAL_CHARS = 120_000
TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".html"}


async def extract_uploaded_files(files: list[UploadFile]) -> tuple[str, list[dict[str, str]]]:
    sections: list[str] = []
    manifest: list[dict[str, str]] = []
    total_chars = 0

    for upload in files:
        name = Path(upload.filename or "file").name
        data = await upload.read()
        if len(data) > MAX_FILE_BYTES:
            raise ValueError(f"الملف {name} يتجاوز الحد الأقصى 10MB")
        suffix = Path(name).suffix.lower()
        text = ""
        status = "مقروء"

        if suffix in TEXT_EXTENSIONS:
            text = data.decode("utf-8", errors="replace")
        elif suffix == ".pdf":
            reader = PdfReader(io.BytesIO(data))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
            if not text.strip():
                status = "PDF بلا طبقة نصية؛ لم يُستخدم OCR"
        elif suffix == ".docx":
            doc = Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    tables.append(" | ".join(cell.text for cell in row.cells))
            if tables:
                text += "\n\n[جداول]\n" + "\n".join(tables)
        else:
            status = "نوع غير مدعوم للاستخراج النصي"

        if text:
            remaining = MAX_TOTAL_CHARS - total_chars
            text = text[:remaining]
            total_chars += len(text)
            sections.append(f"\n--- بداية الملف: {name} ---\n{text}\n--- نهاية الملف: {name} ---")
        manifest.append({"name": name, "status": status, "characters": str(len(text))})
        if total_chars >= MAX_TOTAL_CHARS:
            break

    return "\n".join(sections), manifest
